"""
validate_table.py — Compare generated comparison table against human reference.

Reads both DOCX tables, matches rows using multi-pass algorithm,
and produces a validation report with coverage and accuracy metrics.
Optionally uses Claude CLI for semantic (LLM-based) evaluation.
"""

import json
import logging
import re
import subprocess
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document

log = logging.getLogger(__name__)
PROMPTS_DIR = Path(__file__).parent.parent / "prompts"


def _normalize_text(text: str) -> str:
    """Collapse whitespace and strip."""
    return re.sub(r"\s+", " ", text.strip())


# ---------------------------------------------------------------------------
# 1. DOCX / JSON reading
# ---------------------------------------------------------------------------

def read_docx_table(docx_path: str) -> list:
    """Read first table from a comparison-table DOCX, skip header row."""
    doc = Document(docx_path)
    if not doc.tables:
        raise ValueError(f"No tables found in {docx_path}")

    tbl = doc.tables[0]
    rows = []
    for row in tbl.rows[1:]:  # skip header
        cells = [c.text.strip() for c in row.cells]
        if len(cells) >= 5:
            rows.append({
                "page": cells[0],
                "item": cells[1],
                "previous_version": cells[2],
                "current_version": cells[3],
                "note": cells[4],
            })
    return rows


def read_json_rows(json_path: str) -> list:
    """Read change_rows.json directly."""
    return json.load(open(json_path, encoding="utf-8"))


# ---------------------------------------------------------------------------
# 2. Parsing helpers
# ---------------------------------------------------------------------------

def extract_section_numbers(item_text: str) -> list:
    """
    Extract section number patterns from item text.
    "1. 임상시험계획서 요약\n\n1.1.1. 일차 및 이차 목적" -> ["1", "1.1.1"]
    Also handles "Appendix 1", "Appendix 2" etc.
    """
    numbers = re.findall(r"(?<!\d)(\d+(?:\.\d+)+)(?:\.\s|\s|$)", item_text)
    # Also capture single-digit top-level like "1." at start
    top = re.findall(r"(?:^|\n)\s*(\d+)\.\s", item_text)
    # Handle "Appendix N"
    appendix = re.findall(r"[Aa]ppendix\s+(\d+)", item_text)
    appendix_nums = [f"A{n}" for n in appendix]
    all_nums = list(dict.fromkeys(top + numbers + appendix_nums))
    return all_nums


def _deepest_section_number(item_text: str) -> str:
    """Get the most specific (deepest) section number from item text."""
    nums = extract_section_numbers(item_text)
    if not nums:
        return ""
    # Return the one with most dots (deepest)
    return max(nums, key=lambda n: n.count("."))


def parse_page_range(page_str: str) -> set:
    """
    Parse page string to a set of page numbers.
    "18-35" -> {18..35}, "54, 56, 57" -> {54,56,57}
    "57-60\n67-68" -> {57..60, 67, 68}, "전체" -> sentinel {-1}
    "" -> empty set
    """
    if not page_str or not page_str.strip():
        return set()
    page_str = page_str.strip()
    if page_str == "전체":
        return {-1}  # sentinel

    pages = set()
    # Split by newlines, commas, semicolons
    parts = re.split(r"[,;\n]+", page_str)
    for part in parts:
        part = part.strip()
        m = re.match(r"(\d+)\s*[-–]\s*(\d+)", part)
        if m:
            start, end = int(m.group(1)), int(m.group(2))
            pages.update(range(start, end + 1))
        else:
            m2 = re.match(r"(\d+)", part)
            if m2:
                pages.add(int(m2.group(1)))
    return pages


def compute_similarity(text_a: str, text_b: str) -> float:
    """Fuzzy content similarity using SequenceMatcher."""
    a = _normalize_text(text_a)
    b = _normalize_text(text_b)
    if not a and not b:
        return 1.0
    if not a or not b:
        return 0.0
    return SequenceMatcher(None, a, b).ratio()


def _page_overlap_score(pages_a: set, pages_b: set) -> float:
    """Jaccard-like overlap score for page sets."""
    if not pages_a or not pages_b:
        return 0.0
    # Sentinel match (전체)
    if pages_a == {-1} and pages_b == {-1}:
        return 1.0
    if {-1} & (pages_a | pages_b):
        return 0.5  # partial match with 전체
    intersection = pages_a & pages_b
    union = pages_a | pages_b
    if not union:
        return 0.0
    return len(intersection) / len(union)


def _section_is_child(parent_nums: list, child_nums: list) -> bool:
    """Check if any child section number starts with a parent section number."""
    for cn in child_nums:
        for pn in parent_nums:
            if cn.startswith(pn + ".") or cn == pn:
                return True
    return False


# ---------------------------------------------------------------------------
# 3. Matching algorithm (4-pass)
# ---------------------------------------------------------------------------

def match_rows(ref_rows: list, gen_rows: list) -> list:
    """
    Multi-pass matching: maps each reference row to generated rows.

    Returns list of match dicts (one per reference row), plus excess list.
    """
    # Track which gen rows are used — allow shared usage in pass 3
    # (e.g., ref "header version" and "header date" may both match same gen "header" row)
    used_gen = set()
    match_results = []

    # Pre-compute section numbers and page sets
    ref_info = []
    for r in ref_rows:
        ref_info.append({
            "sections": extract_section_numbers(r["item"]),
            "deepest": _deepest_section_number(r["item"]),
            "pages": parse_page_range(r["page"]),
        })

    gen_info = []
    for g in gen_rows:
        gen_info.append({
            "sections": extract_section_numbers(g["item"]),
            "deepest": _deepest_section_number(g["item"]),
            "pages": parse_page_range(g["page"]),
        })

    # --- Pass 1: Deepest section number exact match ---
    for ri, ref_row in enumerate(ref_rows):
        ref_deepest = ref_info[ri]["deepest"]
        if not ref_deepest:
            match_results.append(None)  # placeholder
            continue

        matched_indices = []
        for gi, gen_row in enumerate(gen_rows):
            if gi in used_gen:
                continue
            gen_deepest = gen_info[gi]["deepest"]
            if gen_deepest == ref_deepest:
                matched_indices.append(gi)

        if matched_indices:
            for gi in matched_indices:
                used_gen.add(gi)
            match_results.append({
                "match_pass": "section_number",
                "matched_gen_indices": matched_indices,
            })
        else:
            match_results.append(None)

    # --- Pass 2: Page overlap + section hierarchy ---
    for ri, ref_row in enumerate(ref_rows):
        if match_results[ri] is not None:
            continue

        ref_pages = ref_info[ri]["pages"]
        ref_sections = ref_info[ri]["sections"]
        if not ref_pages and not ref_sections:
            continue

        matched_indices = []
        for gi, gen_row in enumerate(gen_rows):
            if gi in used_gen:
                continue

            gen_pages = gen_info[gi]["pages"]
            gen_sections = gen_info[gi]["sections"]

            page_overlap = _page_overlap_score(ref_pages, gen_pages) > 0.1
            section_child = _section_is_child(ref_sections, gen_sections) if ref_sections and gen_sections else False

            if page_overlap and section_child:
                matched_indices.append(gi)

        if matched_indices:
            for gi in matched_indices:
                used_gen.add(gi)
            match_results[ri] = {
                "match_pass": "page_section_hierarchy",
                "matched_gen_indices": matched_indices,
            }

    # --- Pass 3: Content similarity + keyword fallback ---
    for ri, ref_row in enumerate(ref_rows):
        if match_results[ri] is not None:
            continue

        ref_pages = ref_info[ri]["pages"]
        ref_item_norm = _normalize_text(ref_row["item"]).lower()
        best_candidates = []

        for gi, gen_row in enumerate(gen_rows):
            if gi in used_gen:
                continue

            gen_item_norm = _normalize_text(gen_row["item"]).lower()

            # Composite score — use max of similarity and containment
            item_sim = compute_similarity(ref_row["item"], gen_row["item"])

            ref_prev = _normalize_text(ref_row.get("previous_version", ""))
            ref_curr = _normalize_text(ref_row.get("current_version", ""))
            gen_prev = _normalize_text(gen_row.get("previous_version", ""))
            gen_curr = _normalize_text(gen_row.get("current_version", ""))

            content_sim = max(
                compute_similarity(ref_prev, gen_prev),
                compute_similarity(ref_curr, gen_curr),
                _containment_score(ref_prev, gen_prev) if ref_prev else 0,
                _containment_score(ref_curr, gen_curr) if ref_curr else 0,
            )
            page_score = _page_overlap_score(ref_pages, gen_info[gi]["pages"])

            # Keyword bonus: check if key terms from ref item appear in gen item/note
            keyword_bonus = 0.0
            gen_note_norm = _normalize_text(gen_row.get("note", "")).lower()
            key_terms = ["머리글", "바닥글", "header", "footer", "표지", "cover",
                         "버전", "날짜", "약어", "abbreviation", "참고문헌",
                         "참고 문헌", "reference", "appendix", "문서 제개정"]
            # Normalize spaces for keyword matching
            ref_item_nospace = ref_item_norm.replace(" ", "")
            gen_item_nospace = gen_item_norm.replace(" ", "")
            gen_note_nospace = gen_note_norm.replace(" ", "")
            for term in key_terms:
                term_nospace = term.replace(" ", "")
                if term_nospace in ref_item_nospace and (term_nospace in gen_item_nospace or term_nospace in gen_note_nospace):
                    keyword_bonus = 0.3
                    break

            # Substring containment bonus
            substring_bonus = 0.0
            if ref_prev and len(ref_prev) >= 2 and ref_prev in gen_prev:
                substring_bonus = 0.3
            if ref_curr and len(ref_curr) >= 2 and ref_curr in gen_curr:
                substring_bonus = max(substring_bonus, 0.3)

            composite = 0.3 * item_sim + 0.5 * content_sim + 0.2 * page_score
            composite += keyword_bonus + substring_bonus
            composite = min(composite, 1.0)

            if composite >= 0.25:
                best_candidates.append((gi, composite))

        if best_candidates:
            best_candidates.sort(key=lambda x: -x[1])
            matched_indices = [best_candidates[0][0]]
            for gi, score in best_candidates[1:]:
                if score >= 0.35:
                    matched_indices.append(gi)
                else:
                    break

            # In Pass 3, allow gen rows to be shared across ref rows
            # (e.g., one "header" gen row covers both "header version"
            # and "header date" ref rows). Only mark used for excess calc.
            pass  # don't add to used_gen here
            match_results[ri] = {
                "match_pass": "content_similarity",
                "matched_gen_indices": matched_indices,
            }

    # --- Pass 4: Fill remaining as missed ---
    for ri in range(len(ref_rows)):
        if match_results[ri] is None:
            match_results[ri] = {
                "match_pass": "none",
                "matched_gen_indices": [],
            }

    return match_results, used_gen


# ---------------------------------------------------------------------------
# 4. Scoring
# ---------------------------------------------------------------------------

def _containment_score(short_text: str, long_text: str) -> float:
    """Check how much of the shorter text is contained in the longer text."""
    s = _normalize_text(short_text).lower()
    l = _normalize_text(long_text).lower()
    if not s or not l:
        return 0.0
    # Check if key phrases from short appear in long
    # Split short into chunks and check containment
    words = s.split()
    if len(words) <= 3:
        return 1.0 if s in l else 0.0
    # Check 5-word windows
    window = min(5, len(words))
    matches = 0
    total = max(1, len(words) - window + 1)
    for i in range(total):
        chunk = " ".join(words[i:i+window])
        if chunk in l:
            matches += 1
    return matches / total


def score_match(ref_row: dict, matched_gen_rows: list) -> dict:
    """Compute content accuracy scores for a matched pair/group."""
    if not matched_gen_rows:
        return {
            "previous_version_sim": 0.0,
            "current_version_sim": 0.0,
            "note_sim": 0.0,
            "combined_score": 0.0,
        }

    # Concatenate generated rows' content
    gen_prev = "\n".join(g.get("previous_version", "") for g in matched_gen_rows)
    gen_curr = "\n".join(g.get("current_version", "") for g in matched_gen_rows)
    gen_note = "\n".join(g.get("note", "") for g in matched_gen_rows)

    ref_prev = ref_row.get("previous_version", "").strip()
    ref_curr = ref_row.get("current_version", "").strip()

    # Use max of SequenceMatcher and containment score
    # (human summaries are often substrings of generated verbatim text)
    prev_sim = max(
        compute_similarity(ref_prev, gen_prev),
        _containment_score(ref_prev, gen_prev),
    )
    curr_sim = max(
        compute_similarity(ref_curr, gen_curr),
        _containment_score(ref_curr, gen_curr),
    )
    note_sim = compute_similarity(ref_row.get("note", ""), gen_note)

    if not ref_prev and not ref_curr:
        # Note-only row (e.g., restructuring changes with empty prev/curr)
        combined = note_sim
    else:
        combined = 0.4 * prev_sim + 0.4 * curr_sim + 0.2 * note_sim

    return {
        "previous_version_sim": round(prev_sim, 3),
        "current_version_sim": round(curr_sim, 3),
        "note_sim": round(note_sim, 3),
        "combined_score": round(combined, 3),
    }


def classify_match(scores: dict) -> str:
    """Classify match status based on combined score."""
    s = scores["combined_score"]
    if s >= 0.6:
        return "full_match"
    elif s >= 0.3:
        return "partial_match"
    elif s > 0:
        return "weak_match"
    return "missed"


# ---------------------------------------------------------------------------
# 5. Report generation
# ---------------------------------------------------------------------------

def generate_report(
    ref_rows: list,
    gen_rows: list,
    match_results: list,
    used_gen: set,
    ref_path: str,
    gen_path: str,
) -> dict:
    """Build the full validation report."""
    matches = []
    missed = []

    # Recompute used_gen from all match results (Pass 3 doesn't add to used_gen)
    all_used_gen = set(used_gen)
    for mr in match_results:
        for gi in mr["matched_gen_indices"]:
            all_used_gen.add(gi)
    used_gen = all_used_gen

    for ri, ref_row in enumerate(ref_rows):
        mr = match_results[ri]
        matched_indices = mr["matched_gen_indices"]
        matched_gen = [gen_rows[gi] for gi in matched_indices]

        scores = score_match(ref_row, matched_gen)

        if not matched_indices:
            status = "missed"
            missed.append({
                "ref_index": ri,
                "ref_row": ref_row,
                "status": "missed",
            })
        else:
            status = classify_match(scores)

        matches.append({
            "ref_index": ri,
            "ref_row": ref_row,
            "status": status,
            "match_pass": mr["match_pass"],
            "matched_gen_indices": matched_indices,
            "matched_gen_count": len(matched_indices),
            "scores": scores,
        })

    # Excess generated rows
    excess = []
    for gi in range(len(gen_rows)):
        if gi not in used_gen:
            excess.append({
                "gen_index": gi,
                "gen_row": gen_rows[gi],
                "status": "excess",
            })

    # Summary metrics
    total_ref = len(ref_rows)
    full = sum(1 for m in matches if m["status"] == "full_match")
    partial = sum(1 for m in matches if m["status"] == "partial_match")
    weak = sum(1 for m in matches if m["status"] == "weak_match")
    missed_count = sum(1 for m in matches if m["status"] == "missed")
    matched_scores = [m["scores"]["combined_score"] for m in matches if m["status"] != "missed"]
    avg_accuracy = round(sum(matched_scores) / len(matched_scores), 3) if matched_scores else 0.0

    summary = {
        "coverage_rate": round((full + partial) / total_ref, 3) if total_ref else 0,
        "full_match_rate": round(full / total_ref, 3) if total_ref else 0,
        "partial_match_rate": round(partial / total_ref, 3) if total_ref else 0,
        "weak_match_rate": round(weak / total_ref, 3) if total_ref else 0,
        "missed_rate": round(missed_count / total_ref, 3) if total_ref else 0,
        "avg_content_accuracy": avg_accuracy,
        "gen_rows_matched": len(used_gen),
        "gen_rows_excess": len(excess),
        "full_match": full,
        "partial_match": partial,
        "weak_match": weak,
        "missed": missed_count,
    }

    return {
        "metadata": {
            "reference_path": ref_path,
            "generated_path": gen_path,
            "timestamp": datetime.now().isoformat(),
            "ref_row_count": len(ref_rows),
            "gen_row_count": len(gen_rows),
        },
        "summary": summary,
        "matches": matches,
        "excess": excess,
    }


def print_summary(report: dict):
    """Print console summary."""
    meta = report["metadata"]
    s = report["summary"]
    total_ref = meta["ref_row_count"]
    total_gen = meta["gen_row_count"]

    print(f"Reference rows:  {total_ref}    Generated rows: {total_gen}")
    print()
    print("--- Match Results ---")
    print(f"  Full match:     {s['full_match']:2d}/{total_ref} ({s['full_match_rate']*100:.1f}%)")
    print(f"  Partial match:  {s['partial_match']:2d}/{total_ref} ({s['partial_match_rate']*100:.1f}%)")
    print(f"  Weak match:     {s['weak_match']:2d}/{total_ref} ({s['weak_match_rate']*100:.1f}%)")
    print(f"  Missed:         {s['missed']:2d}/{total_ref} ({s['missed_rate']*100:.1f}%)")
    print()
    print("--- Coverage ---")
    print(f"  Coverage rate:      {s['coverage_rate']*100:.1f}%  (full + partial)")
    print(f"  Avg content score:  {s['avg_content_accuracy']:.3f}")
    print()
    print("--- Generated Row Analysis ---")
    print(f"  Matched to ref:  {s['gen_rows_matched']:2d}/{total_gen} ({s['gen_rows_matched']/total_gen*100:.1f}%)" if total_gen else "")
    print(f"  Excess (no ref): {s['gen_rows_excess']:2d}/{total_gen} ({s['gen_rows_excess']/total_gen*100:.1f}%)" if total_gen else "")

    # Print missed details
    missed = [m for m in report["matches"] if m["status"] == "missed"]
    if missed:
        print()
        print("--- Missed Reference Rows ---")
        for m in missed:
            r = m["ref_row"]
            print(f"  [{m['ref_index']+1}] page={r['page'][:10]} item={r['item'][:60]}")

    # Print LLM scores if available
    if report.get("llm_summary"):
        ls = report["llm_summary"]
        print()
        print("--- LLM Semantic Evaluation ---")
        print(f"  Full match:     {ls['full_match']:2d}/{total_ref} ({ls['full_match']/total_ref*100:.1f}%)")
        print(f"  Partial match:  {ls['partial_match']:2d}/{total_ref} ({ls['partial_match']/total_ref*100:.1f}%)")
        print(f"  Weak match:     {ls['weak_match']:2d}/{total_ref} ({ls['weak_match']/total_ref*100:.1f}%)")
        print(f"  Missed:         {ls['missed']:2d}/{total_ref} ({ls['missed']/total_ref*100:.1f}%)")
        print(f"  Avg LLM score:  {ls['avg_llm_score']:.3f}")
        print(f"  Coverage (LLM): {ls['coverage_rate']*100:.1f}%")


# ---------------------------------------------------------------------------
# 6. LLM-based semantic evaluation
# ---------------------------------------------------------------------------

def _call_claude_eval(user_prompt: str, system_prompt: str, schema_str: str) -> dict:
    """Call Claude CLI for evaluation and return parsed result."""
    cmd = [
        "claude", "-p", user_prompt,
        "--output-format", "json",
        "--json-schema", schema_str,
        "--append-system-prompt", system_prompt,
    ]
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
    if result.returncode != 0:
        raise RuntimeError(f"Claude CLI failed (exit {result.returncode}): {result.stderr[:300]}")

    response = json.loads(result.stdout)
    structured = response.get("structured_output")
    if not structured:
        raise ValueError("No structured_output in response")
    if isinstance(structured, str):
        structured = json.loads(structured)
    return structured


def _build_eval_prompt(batch: list) -> str:
    """Build evaluation prompt for a batch of matched pairs."""
    parts = []
    for item in batch:
        ri = item["ref_index"]
        ref = item["ref_row"]
        gen_rows = item["gen_rows"]

        parts.append(f"=== Reference Row {ri} ===")
        parts.append(f"Page: {ref.get('page', '')}")
        parts.append(f"Item: {ref.get('item', '')}")
        parts.append(f"Previous: {ref.get('previous_version', '')[:500]}")
        parts.append(f"Current: {ref.get('current_version', '')[:500]}")
        parts.append(f"Note: {ref.get('note', '')}")
        parts.append("")

        for gi, g in enumerate(gen_rows):
            parts.append(f"  --- Generated Row (matched) ---")
            parts.append(f"  Page: {g.get('page', '')}")
            parts.append(f"  Item: {g.get('item', '')}")
            parts.append(f"  Previous: {g.get('previous_version', '')[:500]}")
            parts.append(f"  Current: {g.get('current_version', '')[:500]}")
            parts.append(f"  Note: {g.get('note', '')}")
            parts.append("")

        parts.append("")

    prompt = "\n".join(parts)
    prompt += f"\nEvaluate each of the {len(batch)} reference rows above against their matched generated rows."
    prompt += "\nReturn evaluations for ref_index values: " + ", ".join(str(item["ref_index"]) for item in batch)
    return prompt


def llm_score_matches(report: dict, gen_rows: list, max_workers: int = 3) -> dict:
    """
    Run LLM-based semantic evaluation on matched pairs.
    Updates report in-place with llm_scores per match and llm_summary.
    """
    system_prompt = (PROMPTS_DIR / "validation_eval_system.txt").read_text(encoding="utf-8")
    schema_str = json.dumps(json.load(open(PROMPTS_DIR / "validation_eval_schema.json", encoding="utf-8")))

    # Build evaluation batches (up to 5 pairs per batch)
    BATCH_SIZE = 5
    eval_items = []
    for m in report["matches"]:
        if m["status"] == "missed":
            continue
        eval_items.append({
            "ref_index": m["ref_index"],
            "ref_row": m["ref_row"],
            "gen_rows": [gen_rows[gi] for gi in m["matched_gen_indices"]],
        })

    batches = []
    for i in range(0, len(eval_items), BATCH_SIZE):
        batches.append(eval_items[i:i + BATCH_SIZE])

    log.info(f"LLM evaluation: {len(eval_items)} pairs in {len(batches)} batches")

    # Parallel evaluation
    all_evals = {}  # ref_index -> eval dict
    failed = 0

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_batch = {}
        for batch in batches:
            prompt = _build_eval_prompt(batch)
            future = executor.submit(_call_claude_eval, prompt, system_prompt, schema_str)
            future_to_batch[future] = batch

        for future in as_completed(future_to_batch):
            batch = future_to_batch[future]
            try:
                result = future.result()
                evals = result.get("evaluations", [])
                for ev in evals:
                    ri = ev["ref_index"]
                    all_evals[ri] = {
                        "content_coverage": round(ev.get("content_coverage", 0), 3),
                        "accuracy": round(ev.get("accuracy", 0), 3),
                        "note_quality": round(ev.get("note_quality", 0), 3),
                        "reasoning": ev.get("reasoning", ""),
                        "llm_combined": round(
                            0.5 * ev.get("content_coverage", 0)
                            + 0.3 * ev.get("accuracy", 0)
                            + 0.2 * ev.get("note_quality", 0),
                            3
                        ),
                    }
                log.info(f"  Batch ({len(batch)} pairs) → {len(evals)} evaluations")
            except Exception as e:
                log.warning(f"  Batch failed: {e}")
                failed += 1

    if failed:
        log.warning(f"  Failed batches: {failed}")

    # Update report matches with LLM scores
    for m in report["matches"]:
        ri = m["ref_index"]
        if ri in all_evals:
            m["llm_scores"] = all_evals[ri]
            # Reclassify based on LLM score
            llm_s = all_evals[ri]["llm_combined"]
            if llm_s >= 0.7:
                m["llm_status"] = "full_match"
            elif llm_s >= 0.4:
                m["llm_status"] = "partial_match"
            elif llm_s > 0:
                m["llm_status"] = "weak_match"
            else:
                m["llm_status"] = "missed"
        else:
            m["llm_scores"] = None
            m["llm_status"] = m["status"]  # fallback to text-based

    # Compute LLM summary
    total_ref = len(report["matches"])
    llm_full = sum(1 for m in report["matches"] if m.get("llm_status") == "full_match")
    llm_partial = sum(1 for m in report["matches"] if m.get("llm_status") == "partial_match")
    llm_weak = sum(1 for m in report["matches"] if m.get("llm_status") == "weak_match")
    llm_missed = sum(1 for m in report["matches"] if m.get("llm_status") == "missed")
    llm_scores = [m["llm_scores"]["llm_combined"] for m in report["matches"] if m.get("llm_scores")]
    avg_llm = round(sum(llm_scores) / len(llm_scores), 3) if llm_scores else 0.0

    report["llm_summary"] = {
        "full_match": llm_full,
        "partial_match": llm_partial,
        "weak_match": llm_weak,
        "missed": llm_missed,
        "avg_llm_score": avg_llm,
        "coverage_rate": round((llm_full + llm_partial) / total_ref, 3) if total_ref else 0,
    }

    return report
