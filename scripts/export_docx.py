"""
export_docx.py — Generate comparison table DOCX matching reference format.

Reference format (SAMPLE_PROTOCOL):
- Landscape A4
- 5 columns: Page | Item | Previous Version | Current Version | Note
- Column widths (dxa): 639, 1835, 5296, 5271, 2127
- Header row: Bold, top border single 12pt, bottom border double 4pt
- Data rows: Item column bold only
- Table borders: single 4pt throughout
"""

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Emu, Twips
from docx.enum.section import WD_ORIENT
from lxml import etree


# Reference format dimensions (EMU)
PAGE_WIDTH = 10692130
PAGE_HEIGHT = 7560310
MARGIN_TOP = 1002665
MARGIN_BOTTOM = 900430
MARGIN_LEFT = 720090
MARGIN_RIGHT = 790575

# Column widths in dxa (1 dxa ≈ 635 EMU)
COL_WIDTHS_DXA = [639, 1835, 5296, 5271, 2127]
COL_HEADERS = ["Page", "Item", "Previous Version", "Current Version", "Note"]

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _set_cell_border(cell, **kwargs):
    """
    Set cell borders using lxml direct XML manipulation.
    
    kwargs: top, bottom, left, right, insideH, insideV
    Each value is a dict: {"sz": str, "val": str, "color": str, "space": str}
    """
    tc = cell._tc
    tcPr = tc.find(f"{{{W_NS}}}tcPr")
    if tcPr is None:
        tcPr = etree.SubElement(tc, f"{{{W_NS}}}tcPr")
        tc.insert(0, tcPr)

    tcBorders = tcPr.find(f"{{{W_NS}}}tcBorders")
    if tcBorders is None:
        tcBorders = etree.SubElement(tcPr, f"{{{W_NS}}}tcBorders")

    for edge, attrs in kwargs.items():
        elem = tcBorders.find(f"{{{W_NS}}}{edge}")
        if elem is None:
            elem = etree.SubElement(tcBorders, f"{{{W_NS}}}{edge}")
        for attr, val in attrs.items():
            elem.set(f"{{{W_NS}}}{attr}", val)


def _standard_border():
    """Standard table border: single 4pt."""
    return {"sz": "4", "val": "single", "color": "auto", "space": "0"}


def _header_top_border():
    """Header top border: single 12pt."""
    return {"sz": "12", "val": "single", "color": "auto", "space": "0"}


def _header_bottom_border():
    """Header bottom border: double 4pt."""
    return {"sz": "4", "val": "double", "color": "auto", "space": "0"}


def export_docx(
    rows: list,
    output_path: str,
    old_ver: str = "",
    new_ver: str = "",
):
    """
    Generate comparison table DOCX.
    
    rows: list of dicts with page, item, previous_version, current_version, note
    output_path: path to write DOCX
    old_ver: old version string for header (e.g. "V1.0")
    new_ver: new version string for header (e.g. "V2")
    """
    doc = Document()

    # Page setup: Landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

    # Create table
    num_rows = len(rows) + 1  # +1 for header
    table = doc.add_table(rows=num_rows, cols=5)

    # Set column widths
    for i, width_dxa in enumerate(COL_WIDTHS_DXA):
        for row in table.rows:
            row.cells[i].width = Twips(width_dxa)

    # Header row with version numbers
    headers = COL_HEADERS.copy()
    if old_ver:
        headers[2] = f"Previous Version ({old_ver})"
    if new_ver:
        headers[3] = f"Current Version ({new_ver})"

    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(header_text)
        run.bold = True
        run.font.size = Pt(9)

        # Header borders
        _set_cell_border(
            cell,
            top=_header_top_border(),
            bottom=_header_bottom_border(),
            left=_standard_border(),
            right=_standard_border(),
        )

    # Data rows
    for ri, row_data in enumerate(rows):
        doc_row = table.rows[ri + 1]
        values = [
            row_data.get("page", ""),
            row_data.get("item", ""),
            row_data.get("previous_version", ""),
            row_data.get("current_version", ""),
            row_data.get("note", ""),
        ]
        for ci, value in enumerate(values):
            cell = doc_row.cells[ci]
            para = cell.paragraphs[0]
            para.clear()
            run = para.add_run(str(value))
            run.font.size = Pt(9)
            if ci == 1:  # Item column bold
                run.bold = True

            # Standard borders for data cells
            _set_cell_border(
                cell,
                top=_standard_border(),
                bottom=_standard_border(),
                left=_standard_border(),
                right=_standard_border(),
            )

    # Save
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    print(f"[export_docx] {len(rows)} rows → {output_path}")
    return str(output_path)


def main():
    parser = argparse.ArgumentParser(description="Export comparison table DOCX")
    parser.add_argument("--rows", required=True, help="Path to change_rows.json")
    parser.add_argument("--out", required=True, help="Output DOCX path")
    parser.add_argument("--old-ver", default="", help="Old version string")
    parser.add_argument("--new-ver", default="", help="New version string")
    args = parser.parse_args()

    rows = json.load(open(args.rows, encoding="utf-8"))
    export_docx(rows, args.out, args.old_ver, args.new_ver)


if __name__ == "__main__":
    main()
